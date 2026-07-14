import os
import sys
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Helpers for XML manipulation (cell background color, borders)
def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def main():
    json_path = 'scratch/curriculum_survey_data.json'
    if not os.path.exists(json_path):
        print(f"Error: {json_path} does not exist. Run get_clean_responses.py first.")
        return
        
    with open(json_path, 'r', encoding='utf-8') as f:
        survey_data = json.load(f)
        
    rows = survey_data.get('rows', [])
    
    # Initialize Word Document
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Sleek slate-dark gray
    
    # 1. Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("2026학년도 입학생 3개년 교육과정 편성표 의견 수렴 결과 보고서")
    title_run.font.name = '맑은 고딕'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x03, 0x69, 0xa1) # Premium Blue Accent
    
    # Spacing
    doc.add_paragraph()
    
    # 2. Metadata Table
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = 'Table Grid'
    
    headers = [
        ["일시", "2026. 06. 08.(월)", "대상 교과", "국어과 (전체 교사)"],
        ["대상 인원", "총 8명", "참여 인원", "8명 (참여율 100%)"]
    ]
    
    for r_idx, row_data in enumerate(headers):
        for c_idx, text in enumerate(row_data):
            cell = meta_table.rows[r_idx].cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9.5)
            
            # Formatting keys and values
            if c_idx % 2 == 0:
                run.font.bold = True
                set_cell_background(cell, "F1F5F9") # slate-100 background
                
    doc.add_paragraph() # Spacing
    
    # 3. Section 1: 투표 결과
    h1 = doc.add_heading(level=2)
    h1_run = h1.add_run("1. 안별 최종 투표 현황")
    h1_run.font.name = '맑은 고딕'
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x03, 0x69, 0xa1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    # Calculate stats
    total_votes = len(rows)
    opt1_votes = sum(1 for r in rows if len(r) > 2 and r[2] == '1안')
    opt2_votes = sum(1 for r in rows if len(r) > 2 and r[2] == '2안')
    opt3_votes = sum(1 for r in rows if len(r) > 2 and r[2] == '3안')
    opt4_votes = sum(1 for r in rows if len(r) > 2 and r[2] == '4안')
    
    stats_table = doc.add_table(rows=5, cols=3)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_table.style = 'Table Grid'
    
    stats_data = [
        ["선택안", "득표수", "득표율 (%)"],
        ["1안 (이전과 동일 / 일부 과목 제외)", f"{opt1_votes}표", f"{(opt1_votes/total_votes)*100:.1f}%"],
        ["2안 (선택군 C-D 통합)", f"{opt2_votes}표", f"{(opt2_votes/total_votes)*100:.1f}%"],
        ["3안 (C-D / E-F 각 통합)", f"{opt3_votes}표", f"{(opt3_votes/total_votes)*100:.1f}%"],
        ["4안 (지정 축소 및 E-F 통합)", f"{opt4_votes}표", f"{(opt4_votes/total_votes)*100:.1f}%"]
    ]
    
    for r_idx, row_data in enumerate(stats_data):
        for c_idx, text in enumerate(row_data):
            cell = stats_table.rows[r_idx].cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            p = cell.paragraphs[0]
            if c_idx > 0 or r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p.runs[0]
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
            
            if r_idx == 0:
                run.font.bold = True
                set_cell_background(cell, "E0F2FE") # Sky blue background
                
    doc.add_paragraph() # Spacing
    
    # 4. Section 2: 교사별 의견 종합표
    h2 = doc.add_heading(level=2)
    h2_run = h2.add_run("2. 교사별 선택안 및 상세 의견")
    h2_run.font.name = '맑은 고딕'
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x03, 0x69, 0xa1)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    survey_table = doc.add_table(rows=total_votes + 1, cols=4)
    survey_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    survey_table.style = 'Table Grid'
    
    # Table headers
    headers_survey = ["제출 교사", "선택 안", "상세 의견 및 건의사항", "제출일시"]
    for c_idx, text in enumerate(headers_survey):
        cell = survey_table.rows[0].cells[c_idx]
        cell.text = text
        set_cell_background(cell, "E2E8F0") # Slate-200 background
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = '맑은 고딕'
        run.font.bold = True
        run.font.size = Pt(9.5)
        
    # Table body
    for r_idx, row_data in enumerate(rows):
        cells = survey_table.rows[r_idx + 1].cells
        
        # Name
        name = row_data[1]
        # Option
        opt = row_data[2]
        # Time
        timestamp = row_data[0]
        # Opinion extraction
        opinions = []
        if len(row_data) > 3 and row_data[3] and row_data[3].strip() != ".":
            opinions.append(f"1안 의견: {row_data[3].strip()}")
        if len(row_data) > 5 and row_data[5] and row_data[5].strip() != ".":
            opinions.append(f"2안/3안 의견: {row_data[5].strip()}")
        if len(row_data) > 7 and row_data[7] and row_data[7].strip() != ".":
            opinions.append(f"2안/3안/4안 의견: {row_data[7].strip()}")
            
        opinion_text = "\n".join(opinions) if opinions else "(특이 의견 없음)"
        
        data_row = [name, opt, opinion_text, timestamp]
        for c_idx, text in enumerate(data_row):
            cell = cells[c_idx]
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            
            if c_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            run = p.runs[0]
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            if c_idx == 1:
                run.font.bold = True
                if text == '1안':
                    run.font.color.rgb = RGBColor(0x02, 0x84, 0xc7) # Sky-600
                elif text == '2안':
                    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Emerald-600
                elif text == '3안':
                    run.font.color.rgb = RGBColor(0xd9, 0x77, 0x06) # Amber-600
                    
    doc.add_paragraph() # Spacing
    
    # 5. Section 3: 주요 쟁점 및 건의사항
    h3 = doc.add_heading(level=2)
    h3_run = h3.add_run("3. 교과 의견별 쟁점 요약")
    h3_run.font.name = '맑은 고딕'
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x03, 0x69, 0xa1)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("■ 1안 지지 논지 (37.5%, 3표)")
    p1 = doc.add_paragraph(" - 국어 교과의 수업 시수 확보와 교육과정 경쟁력 유지가 주된 이유입니다.\n"
                           " - 학생들에게 선택지가 너무 넓어지면, 과목 분산 선택으로 인한 국어과 폐강 위험이 있으며 교사들의 다과목 담당 행정 및 평가 부담이 가중될 것을 지적했습니다.")
    p1.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph("■ 2안 지지 논지 (37.5%, 3표)")
    p2 = doc.add_paragraph(" - 지지하는 세 교사 모두 주관식 코멘트는 적지 않았으나, 1안의 타이트한 설계보다는 학생들의 기초적인 선택권 확대(선택군 C-D 통합)를 원하는 성향을 보이고 있습니다.")
    p2.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph("■ 3안 지지 논지 (25.0%, 2표)")
    p3 = doc.add_paragraph(" - 국어과 시수를 안전하게 확보하면서도 학생들의 선택권을 유연하게 넓힐 수 있는 양립(상생) 절충안으로서 지지하는 의견을 보였습니다.")
    p3.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_paragraph()
    
    # 6. Section 4: 결론 및 제안
    h4 = doc.add_heading(level=2)
    h4_run = h4.add_run("4. 결론 및 향후 절차 제안")
    h4_run.font.name = '맑은 고딕'
    h4_run.font.size = Pt(13)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0x03, 0x69, 0xa1)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("1) 삼파전 대립 구도 형성\n"
                      "   대상 교사 8명 전원의 참여 결과 1안(3표), 2안(3표), 3안(2표)으로 표가 분산되어 확실한 단일 안이 도출되지 못했습니다.")
    doc.add_paragraph("2) 대면 협의회 개최 권장\n"
                      "   1안의 '교사 부담 경감 및 시수 사수' 논리와 2안의 '선택권 확대' 논리가 팽팽히 맞서고 있습니다. 이에 따라 양자의 입장을 절충해 낸 3안(시수 확보 + 유연성 제공 모델)을 주제로 삼아 대면 교과 협의회를 개최하여 의견을 하나로 모으는 과정을 거치기를 권장합니다.")
                      
    # Save the file in the workspace
    output_filename = "2026학년도_교육과정_편성표_의견수렴_결과보고서.docx"
    output_path = os.path.join(os.getcwd(), output_filename)
    doc.save(output_path)
    
    print(f"Report generated successfully: {output_path}")

if __name__ == '__main__':
    main()
