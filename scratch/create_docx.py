import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_table_doc():
    doc = Document()
    
    # Title
    heading = doc.add_heading('2. 수업 나눔 계획', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacing
    
    # Table data
    data = [
        ["교사명", "일시", "장소", "대상학급", "과목", "단원 또는 수업 주제", "수업 형태"],
        ["강지영", "2026. 5. 14.(목) 3교시", "2-7", "2-7", "문학", "3. 역사와 함께 흐르는 문학 (상고 시대~고려 시대 문학, 조선 시대 문학)", "작품 감상 및 비평"],
        ["강필성", "2026. 5. 14.(목) 3교시", "2-5", "2-5", "문학", "2. 다양한 빛깔로 만나는 문학 (02. 이생규장전)", "모둠 토의 및 질문수업"],
        ["김승우", "2026. 5. 14.(금) 4교시", "1-5", "1-5", "공통국어1", "1. 네 가지 빛깔로 만나는 문학 (02. 서사 갈래)", "강의 및 내용 분석"],
        ["김태영", "2026. 5. 14.(목) 3교시", "1-9", "1-9", "공통국어1", "4. 세상을 밝히는 논증의 기술 (02. 토론으로 해결방안 탐색하기)", "찬반 토론 및 실습"],
        ["이병의", "[원본 확인 요망]", "[확인 요망]", "[확인 요망]", "문학", "3. 역사와 함께 흐르는 문학 (04. 개화기~일제강점기 문학 / 작품3 천변 풍경)", "내용 분석 및 발표"],
        ["임언숙", "2026. 5. 14.(목) 4교시", "3-1", "3-1", "화법과 작문", "Ⅲ. 현대소설 (01. 소설가 구보씨의 일일)", "모둠 탐구 및 글쓰기"],
        ["조진희", "2026. 5. 14.(목) 4교시", "3-5", "3-5", "화법과 작문", "Ⅱ. 적용학습 (1) 고전시가 – 시조 5편", "짝 활동 및 비교 분석"],
        ["황요한", "2026. 5. 14.(금) 4교시", "3-8", "3-8", "화법과 작문", "수능특강 문학 갈래복합 05 (임유휴, 「목동가」)", "모둠 토의 및 에세이 작성"]
    ]
    
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_data
            # Basic formatting
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    if row_idx == 0:
                        run.font.bold = True
    
    # Save to desktop
    desktop_path = r"D:\OneDrive - 경상남도교육청\바탕 화면"
    output_path = os.path.join(desktop_path, "수업_나눔_계획_표.docx")
    
    doc.save(output_path)
    print(f"Saved to: {output_path}")

if __name__ == "__main__":
    create_table_doc()
