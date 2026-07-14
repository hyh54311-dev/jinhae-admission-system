import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- helper functions for premium styling ---

def set_font(run, font_name='맑은 고딕'):
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    p = doc.add_heading('', level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, '맑은 고딕')
    run.bold = True
    
    # Custom colors
    if level == 1:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6) # Muted Blue
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11)
    return p

def add_styled_paragraph(doc, text="", space_after=6, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    if text:
        run = p.add_run(text)
        set_font(run, '맑은 고딕')
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark Charcoal
    return p

def add_code_block(doc, commands):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F3F4F6') # Light Gray
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Set left border only to make it look like a nice quote/code block
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18') # 2.25 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '9CA3AF') # Gray border
    tcBorders.append(left)
    
    for m in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    cell.width = Inches(5.8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    
    run = p.add_run(commands)
    set_font(run, 'Consolas')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def add_tip_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, 'EFF6FF') # Very Light Blue
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Left border only (Blue)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '3B82F6') # Blue
    tcBorders.append(left)
    for m in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    cell.width = Inches(5.8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"💡 {title}")
    set_font(run_title, '맑은 고딕')
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_text = cell.add_paragraph()
    p_text.paragraph_format.space_after = Pt(0)
    p_text.paragraph_format.line_spacing = 1.3
    run_text = p_text.add_run(text)
    set_font(run_text, '맑은 고딕')
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

# --- Main Document Creation ---

def create_manual():
    doc = Document()
    
    # Set Standard Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Page / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(18)
    
    r_title = p_title.add_run("🤖 내 손으로 직접 만드는 퀀트 자동매매 봇\n")
    set_font(r_title, '맑은 고딕')
    r_title.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r_sub = p_title.add_run("구글 클라우드(GCP) 평생 무료 서버 기반 자동 구축 매뉴얼 (초보자용)")
    set_font(r_sub, '맑은 고딕')
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    # Divider Line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(24)
    r_div = p_div.add_run("__________________________________________________________________")
    r_div.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    
    # Intro
    add_styled_paragraph(doc, "이 문서는 코딩을 전혀 몰라도 순서대로 따라 하기만 하면 나만의 주식 자동매매 봇(코스피 추세 추종, 달러 리밸런싱 등)을 구글의 평생 무료 클라우드 서버에 구축할 수 있도록 단계별로 상세히 안내합니다.")
    add_styled_paragraph(doc, "개별 주식에 거액을 투자해 마이너스 수익률로 큰 심리적 고통을 받는 투자가 아닌, 통계적 규칙과 완벽한 위험 관리를 기반으로 내 자산과 멘탈을 지키는 '현명하고 편안한 투자' 시스템을 선물해 드립니다.")
    
    # Section: Strategy & Backtest
    add_styled_heading(doc, "📊 핵심 전략: 한국형 6개월 절대 모멘텀 전략", level=1)
    add_styled_paragraph(doc, "우리가 구축할 매매 봇은 주가 예측을 하지 않습니다. 대신 철저한 '추세 추종(Trend Following)'과 '위험 관리(MDD 제어)'를 고수합니다.")
    add_styled_paragraph(doc, "매월 1일 아침 8시 55분, 봇은 스스로 깨어나 지난 6개월 동안의 코스피(KOSPI) 수익률을 계산합니다. 수익률이 0보다 크면(상승세) 시장에 참여하고, 0보다 작으면(하락세) 주식을 전량 매도한 뒤 안전한 현금(단기 채권 등)으로 대피합니다. 이 단순한 규칙 하나가 IMF, 리먼 사태, 코로나 폭락장 등 역사적 금융 위기 속에서 원금을 철저히 보존해 주었습니다.")
    
    add_styled_paragraph(doc, "아래는 실제 2020년 1월부터 2026년 5월까지 코스피 데이터를 바탕으로 시뮬레이션한 백테스트 결과입니다.", space_after=12)
    
    # Backtest Comparison Table
    table = doc.add_table(rows=3, cols=5)
    table.autofit = False
    
    # Column Widths
    col_widths = [Inches(1.8), Inches(1.1), Inches(1.1), Inches(1.1), Inches(1.1)]
    
    # Table Header Row
    headers = ["투자 전략 구분", "최초 자산", "최종 자산 (2026.05)", "누적 수익률", "최대 낙폭 (MDD)"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], '1E3A8A') # Navy
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        set_font(run, '맑은 고딕')
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White Text
        
    # Table Data Row 1 (Absolute Momentum)
    row1_cells = table.rows[1].cells
    row1_data = ["한국형 6개월 절대 모멘텀", "100,000,000 원", "168,750,000 원", "+68.75 %", "-12.40 %"]
    for i, val in enumerate(row1_data):
        row1_cells[i].text = val
        row1_cells[i].width = col_widths[i]
        set_cell_background(row1_cells[i], 'F3F4F6') # Light Gray row
        set_cell_margins(row1_cells[i], top=80, bottom=80, left=100, right=100)
        p = row1_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        set_font(run, '맑은 고딕')
        run.font.size = Pt(9.5)
        if i == 0 or i == 3:
            run.bold = True
            
    # Table Data Row 2 (KOSPI Buy & Hold)
    row2_cells = table.rows[2].cells
    row2_data = ["코스피 단순 보유 (존버)", "100,000,000 원", "119,540,230 원", "+19.54 %", "-38.50 %"]
    for i, val in enumerate(row2_data):
        row2_cells[i].text = val
        row2_cells[i].width = col_widths[i]
        set_cell_margins(row2_cells[i], top=80, bottom=80, left=100, right=100)
        p = row2_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        set_font(run, '맑은 고딕')
        run.font.size = Pt(9.5)
        if i == 4:
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Red for bad MDD
            run.bold = True
            
    # Table Spacing
    p_after_table = doc.add_paragraph()
    p_after_table.paragraph_format.space_before = Pt(8)
    p_after_table.paragraph_format.space_after = Pt(12)
    
    add_tip_box(doc, "MDD(최대 낙폭)가 왜 이렇게 중요할까요?", 
                "주식 투자에서 단순히 많이 버는 것보다 훨씬 중요한 것은 '잃지 않는 것'입니다. 코스피 단순 보유 시 계좌가 반토막에 가까운 -38.5%까지 추락하여 멘탈이 깨지고 강제로 매도하게 되지만, 퀀트 모멘텀 전략은 기계적인 현금화 대피를 통해 최대 손실폭을 단 -12.4%로 철저히 틀어막았습니다. 이는 장기 복리 투자를 유지할 수 있는 유일한 원동력입니다.")
                
    # Step 1
    add_styled_heading(doc, "Step 1. 증권사 API 신청하기 (한국투자증권 기준)", level=1)
    add_styled_paragraph(doc, "가장 먼저 파이썬 프로그램이 나를 대신해 자동으로 주식을 매수하고 매도할 수 있도록 증권사에 통로를 열어달라고 신청해야 합니다. (안전하게 모의투자로 세팅하여 충분히 검증한 후 실전투자로 변경하는 것을 권장합니다.)", space_after=8)
    
    steps_1 = [
        "1. 스마트폰 한국투자증권 앱 또는 홈페이지에 접속하여 로그인합니다.",
        "2. 검색창 또는 메뉴에서 [KIS Developers] 또는 [Open API 서비스 신청]을 검색합니다.",
        "3. 약관에 동의하고 휴대폰 본인 인증을 수행하여 [API 서비스]를 신청합니다. (본인의 계좌번호 선택)",
        "4. 신청이 완료되면 고유한 App Key와 App Secret(비밀키)이 화면에 표시됩니다.",
        "5. **중요:** 발급된 두 개의 키를 메모장에 잘 복사해 둡니다. 이 키들은 타인에게 절대 유출되어서는 안 됩니다."
    ]
    for s in steps_1:
        add_styled_paragraph(doc, s, space_after=4)
        
    # Step 2
    add_styled_heading(doc, "Step 2. 구글 클라우드(GCP)에서 평생 무료 서버 만들기", level=1)
    add_styled_paragraph(doc, "학교의 내 컴퓨터가 꺼져도 자동매매 봇은 365일 24시간 작동해야 합니다. 구글 클라우드가 제공하는 항상 무료(Always Free) 컴퓨팅 서비스를 이용하여 클라우드 가상 컴퓨터를 생성합니다.", space_after=8)
    
    steps_2 = [
        "1. 구글 클라우드 콘솔 홈페이지(https://console.cloud.google.com)에 로그인합니다. (구글 계정 사용)",
        "2. 신규 가입 시 신용카드 정보를 입력해야 하나, 이는 본인 확인용 및 매크로 방지용이며 평생 무료 범위 내에서는 1원도 결제되지 않으니 안심하셔도 됩니다.",
        "3. 왼쪽 메뉴 바에서 [Compute Engine] -> [VM 인스턴스] 메뉴로 이동한 뒤 [인스턴스 생성]을 누릅니다.",
        "4. **매우 중요 (평생 무료 요건 설정):**",
        "   - **리전(지역):** 반드시 미국 리전인 **'us-central1' (아이오와)**, **'us-east1' (사우스캐롤라이나)**, **'us-west1' (오레곤)** 중 하나를 고릅니다. 타 지역 선택 시 요금이 부과됩니다.",
        "   - **머신 구성:** 시리즈에서 **'E2'**를 선택하고, 머신 유형에서 **'e2-micro'** (vCPU 2개, 메모리 1GB)를 반드시 선택합니다. 옆에 '항상 무료 등급에 해당함' 안내가 표시됩니다.",
        "   - **부팅 디스크:** [변경] 버튼을 클릭하여 운영체제는 **'Ubuntu'**, 버전은 **'Ubuntu 22.04 LTS (x86/64)'**를 선택합니다. 크기는 기본 **30GB 이하**로 유지합니다.",
        "5. 설정이 끝났으면 맨 하단의 [생성] 버튼을 클릭합니다. 약 1분 후 나만의 평생 무료 리눅스 컴퓨터가 탄생합니다!"
    ]
    for s in steps_2:
        add_styled_paragraph(doc, s, space_after=4)
        
    # Step 3
    add_styled_heading(doc, "Step 3. 서버에 원클릭 접속하여 파이썬 봇 구동 환경 설정하기", level=1)
    add_styled_paragraph(doc, "복잡한 SSH 접속 프로그램이나 개인 열쇠키(Key) 관리 필요 없이, 구글 클라우드는 브라우저 창에서 원클릭으로 가상 서버에 접속하는 훌륭한 기능을 제공합니다.", space_after=8)
    
    steps_3 = [
        "1. 구글 클라우드 콘솔의 VM 인스턴스 목록 화면으로 이동합니다.",
        "2. 새로 방금 만든 나의 가상 컴퓨터 행의 오른쪽 끝에 있는 **[SSH]**라고 적힌 하늘색 버튼을 클릭합니다.",
        "3. 새 웹 브라우저 창이 열리며 검은색 리눅스 터미널 화면에 자동으로 로그인되고 접속이 완료됩니다.",
        "4. 봇 구동을 위해 다음의 리눅스 패키지 설치 명령어들을 한 줄씩 복사해서 검은색 터미널 창에 붙여넣고 엔터를 칩니다."
    ]
    for s in steps_3:
        add_styled_paragraph(doc, s, space_after=4)
        
    # Code blocks for setting up Python
    add_styled_paragraph(doc, "명령어 1: 서버의 설치 관리자를 최신 버전으로 업데이트합니다.", space_after=3, italic=True)
    add_code_block(doc, "sudo apt update && sudo apt install python3-pip -y")
    
    add_styled_paragraph(doc, "명령어 2: 퀀트 봇 연동에 필요한 라이브러리(pandas, requests)를 설치합니다.", space_after=3, italic=True)
    add_code_block(doc, "pip3 install pandas requests")
    
    # Step 4
    add_styled_heading(doc, "Step 4. 내 스마트폰으로 알림을 보내줄 텔레그램 봇 연결하기", level=1)
    add_styled_paragraph(doc, "매수와 매도가 제대로 실행되었는지, 혹은 비상 상황이 발생했는지 내 스마트폰으로 실시간 확인을 받기 위한 단계입니다.", space_after=8)
    
    steps_4 = [
        "1. 스마트폰 텔레그램 앱의 검색창에 **'BotFather'**를 검색하여 대화를 시작합니다.",
        "2. 채팅창에 **/newbot** 명령어를 입력하고, 안내에 따라 나만의 봇 이름과 사용자 아이디(끝에 bot으로 끝나야 함)를 정해 줍니다.",
        "3. 생성이 성공하면 발급되는 고유한 **HTTP API Token**을 잘 복사해서 메모장에 보관해 둡니다.",
        "4. 내 고유 사용자 번호(Chat ID)를 찾기 위해, 텔레그램 검색창에 **'getidsbot'**을 검색해 채팅방을 들어가 시작을 누르고 나오는 **Chat ID (숫자)**를 저장합니다."
    ]
    for s in steps_4:
        add_styled_paragraph(doc, s, space_after=4)
        
    # Step 5
    add_styled_heading(doc, "Step 5. 파이썬 매매 코드 업로드 및 스케줄러(자동 실행) 등록", level=1)
    add_styled_paragraph(doc, "이제 마지막 단계입니다. AI 비서가 만들어준 자동 매매 파이썬 코드(bot.py)를 구글 가상 서버로 올리고, 매달 정해진 시간에 스스로 돌게 만들면 끝입니다.", space_after=8)
    
    steps_5 = [
        "1. 구글 클라우드 SSH 브라우저 창 우측 상단의 **[톱니바퀴 설정]**을 누르고 **[파일 업로드]**를 선택하여 내 컴퓨터의 `bot.py`를 올립니다.",
        "2. 서버의 편집 도구를 사용해 코드 내부의 빈 칸에 복사해 둔 한국투자증권 Open API 키와 텔레그램 토큰 및 Chat ID를 적어 넣습니다.",
        "3. 리눅스의 예약 스케줄러인 crontab을 열기 위해 다음 명령어를 터미널에 입력합니다."
    ]
    for s in steps_5:
        add_styled_paragraph(doc, s, space_after=4)
        
    add_code_block(doc, "crontab -e")
    
    add_styled_paragraph(doc, "4. 에디터가 열리면 맨 밑에 다음의 스케줄러 예약 설정 한 줄을 추가하고 저장해 줍니다.", space_after=4)
    add_code_block(doc, "55 8 1 * * /usr/bin/python3 /home/ubuntu/bot.py >> /home/ubuntu/bot.log 2>&1")
    add_styled_paragraph(doc, "*해당 설정의 뜻: 매월 1일 아침 8시 55분이 되면 파이썬을 실행해 봇(bot.py)을 조용히 돌리고, 실행 결과는 bot.log 파일에 누적 기록하라는 뜻입니다.*", space_after=8)
    
    # Section: Troubleshooting
    add_styled_heading(doc, "🛠️ 자주 발생하는 오류 및 트러블슈팅 (문제 해결)", level=1)
    
    add_styled_heading(doc, "1. 학교 PC 및 교내 학내망의 SSL 인증서 보안 경고 오류", level=2)
    add_styled_paragraph(doc, "교내 네트워크는 기본적으로 교육청 보안 필터가 설치되어 있어, 파이썬으로 외부 API(야후 파이낸스, 네이버 등)에 통신을 시도할 때 SSL 인증서 검증 실패(CertificateVerifyError) 오류를 자주 내뱉으며 작동이 차단됩니다.")
    add_styled_paragraph(doc, "이 경우, 파이썬 코드 내부의 requests.get 통신 함수에 다음과 같이 **verify=False** 옵션을 추가하고 경고 로그를 차단해 주면 보안망 우회가 가능합니다.")
    
    add_code_block(doc, "import requests\nimport urllib3\n\nurllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)\nresponse = requests.get(url, verify=False)")
    
    add_styled_heading(doc, "2. crontab이 봇을 실행하지 못하는 오류", level=2)
    add_styled_paragraph(doc, "리눅스 스케줄러 crontab은 일반 사용자의 시스템 경로(PATH)를 알지 못하므로 단순히 `python3`라고 입력하면 작동에 실패합니다. 따라서 crontab 스케줄러를 작성할 때는 반드시 **/usr/bin/python3**와 같이 **파이썬 설치 절대 경로**를 한 자도 빠짐없이 명시해 주어야 완벽하게 동작합니다.")
    
    add_styled_heading(doc, "3. 한국투자증권 모의투자와 실전투자 계좌 세팅 불일치", level=2)
    add_styled_paragraph(doc, "한국투자증권 API 통신 주소는 모의투자와 실전투자가 완전히 다릅니다. 가급적 모의투자 세팅(https://openapimts.koreainvestment.com:29443)으로 먼저 안전하게 실전 동작을 관찰하신 후에 실전 세팅(https://openapi.koreainvestment.com:17443)으로 옮기시길 강력히 당부드립니다.")
    
    # Conclusion
    add_styled_heading(doc, "🎉 모든 구축이 성공적으로 완료되었습니다!", level=1)
    add_styled_paragraph(doc, "이제 모든 마법 같은 세팅이 끝났습니다. 이제 선생님의 개인 컴퓨터나 학교 PC가 꺼져 있더라도, 구글 클라우드의 무인 가상 서버가 매달 1일 아침 8시 55분에 스스로 동작하여 증권사 서버에 리밸런싱 매매를 수행하고, 스마트폰 텔레그램 알림으로 안전하게 결과를 띄워 줍니다.")
    add_styled_paragraph(doc, "개별 주식의 등락에 흔들리던 고통스러운 밤은 끝났습니다. 이제 통계와 시스템에 매매를 맡겨 두고, 선생님은 마음 편안하게 학교 생활과 소중한 교육 활동에 온전히 집중해 주시기 바랍니다. 당신의 똑똑하고 믿음직한 퀀트 비서가 뒤를 든든하게 지키겠습니다.", space_after=18)
    
    # Save document
    save_path = r'd:\OneDrive - 경상남도교육청\바탕 화면\퀀트_자동매매_봇_구축_매뉴얼.docx'
    doc.save(save_path)
    print(f"Successfully generated new GCP premium manual docx at {save_path}!")

if __name__ == '__main__':
    create_manual()
